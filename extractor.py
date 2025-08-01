# =============================== Imports ===============================
import os
import io
import re
import sys
import fitz
import glob
import shutil
import difflib
import logging
import pdfplumber
import pytesseract
import pandas as pd
import tkinter as tk
from PIL import Image
from docx import Document
from datetime import datetime
from tkinter import filedialog
from bs4 import BeautifulSoup
import google.generativeai as genai

# =============================== Configuration ===============================
GEMINI_API_KEYS = [
    # Add your API keys here
    "AIzaSyCzAr33J_LQPKYNz8L7iqWTL48FKBLRTg0",
    "AIzaSyBJ5sQL299zcDkewRsD9nGbBEIez1oEnYA",
    "AIzaSyB-nuSLbOVrt22WfgNkdAzKjBA4dARRRv4",
    "AIzaSyANJ_9pPk9Y2Hs0rSIVyNNje4aFn6WTi1I",
    "AIzaSyBBz_aYJ3WH5aiG7vNp93wJDSfBJu7OBz0",
    "AIzaSyBzzx2_CqVP-JmUBRgKIvHEkDemGiRuryY",
    "AIzaSyDgqd9_jG-PQT7R1eus3TEorlbRgdnDDog",
    "AIzaSyC4fDaQO7mEsjkXsdRjwnZy_0dCUa3m1UQ",
    "AIzaSyAAU0YRbA-6TEBQGInQoG-W4mlpFUC_Flk",
    "AIzaSyAuxm_KGrmG4BQwP0U8DCqWh0tZuco1tAs",
    "AIzaSyBrDyXjwAHygNwxm_TkyXiDmnYIAeKINjw",
    "AIzaSyA7B_4UUEWqOLrWUwG8APwRIorkpUsjNLc",
    "AIzaSyAtwgN8kAVuutAVEI5pBlpuWSWZvnXXvuM",
    "AIzaSyBPPnzAyYjwGoSIrLoVzVHcB9-b8Rrlv4g",
    "AIzaSyANpCbzw2piGsnaU8YG5Bp-NaCEAjM_e2E",
    # ... (other keys)
]
current_key_index = 0

# Get Pytesseract Path for OCR
def get_tesseract_path():
    """Find or ask for Tesseract path, and save it for future use."""
    saved_path_file = os.path.join(os.path.expanduser("~"), ".tesseract_path.txt")
    # Try to read saved path
    if os.path.exists(saved_path_file):
        with open(saved_path_file, "r") as f:
            saved_path = f.read().strip()
            if os.path.exists(saved_path):
                return saved_path
    # Try auto-detect from PATH
    auto_path = shutil.which("tesseract")
    if auto_path:
        return auto_path
    # Ask user to locate tesseract.exe
    print("🔍 Tesseract not found. Please select your Tesseract executable.")
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title="Select tesseract.exe", filetypes=[("Tesseract", "tesseract.exe")])
    if file_path:
        with open(saved_path_file, "w") as f:
            f.write(file_path)
        return file_path
    raise FileNotFoundError("Tesseract executable not found.")

pytesseract.pytesseract.tesseract_cmd = get_tesseract_path()

# Get Resources Path
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller."""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

INPUT_RULES_FILE = resource_path("Input_Data.xlsx")
MODEL_NAME = "gemini-1.5-flash"  # Use the latest model name

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Configure Gemini API Key
def configure_gemini_api(index):
    """Configure Gemini API with the given key index."""
    key = GEMINI_API_KEYS[index]
    genai.configure(api_key=key)
    logging.info(f"Switched to API Key #{index + 1}")
    return genai.GenerativeModel(MODEL_NAME)

model = configure_gemini_api(current_key_index)

# Standardize Titles
def normalize_title(title):
    """Standardize title by removing brackets and lowering case."""
    title = str(title).strip().lower()
    if title.startswith("[") and title.endswith("]"):
        title = title[1:-1].strip()
    return title

# =============================== Extraction Functions ===============================

# Extract text from PDF files
def extract_text_pages_from_pdf(path):
    """Extract text from each page of a PDF file, using OCR for image-based pages, and extract tables with marks."""
    doc = fitz.open(path)
    pages = []
    for i in range(doc.page_count):
        page = doc.load_page(i)
        text = page.get_text().strip()
        if text:
            pages.append(f"[Page {i+1}]\n{text}")
        else:
            # OCR fallback for scanned pages
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_data))
            ocr_text = pytesseract.image_to_string(image).strip()
            ocr_text = normalize_ocr_ticks(ocr_text)
            ocr_text = postprocess_ocr_text(ocr_text)
            if ocr_text:
                pages.append(f"[Page {i+1} - OCR]\n{ocr_text}")
    # Extract tables with marks
    tables = extract_tables_with_marks_from_pdf(path)
    pages.extend(tables)
    return pages

# Normalize OCR ticks
def normalize_ocr_ticks(text):
    """Convert common OCR tick mark misreads like 'v', 'V', 'Y', etc. into [TICKED] where appropriate."""
    lines = text.splitlines()
    corrected_lines = []
    for line in lines:
        # Heuristic: If line has only 'v', 'V', 'Y' or similar — likely a checkmark
        if re.fullmatch(r"[\s\.\-]*[vVyY][\s\.\-]*", line.strip()):
            corrected_lines.append("[TICKED]")
        else:
            # Replace inline " v ", " V ", etc. with [TICKED] if near known patterns
            fixed = re.sub(r"\b[vVYy]\b", "[TICKED]", line)
            corrected_lines.append(fixed)
    return "\n".join(corrected_lines)

# Postprocess OCR text
def postprocess_ocr_text(text):
    """
    Group OCR lines together that look like table rows (especially with [TICKED]),
    so Gemini understands context and selection logic.
    """
    lines = text.splitlines()
    output = []
    current_table = []
    for line in lines:
        # Likely a table-like row if it has multiple numbers or ticked values
        if any(sym in line for sym in ["|", "[TICKED]"]) or re.search(r"\d+\.\d+", line):
            current_table.append(line.strip())
        else:
            if current_table:
                output.append("[Table]")
                output.extend(current_table)
                output.append("[/Table]")
                current_table = []
            output.append(line.strip())
    if current_table:
        output.append("[Table]")
        output.extend(current_table)
        output.append("[/Table]")
    return "\n".join(output)

# Extract tables from PDF files and annotate ticked/checked values
def extract_tables_with_marks_from_pdf(path):
    """Extract tables from PDF and annotate ticked/checked values."""
    tables = []
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            for table in page.extract_tables():
                processed_rows = []
                for row in table:
                    if not any(row):  # skip empty rows
                        continue
                    processed_row = []
                    for cell in row:
                        cell_str = str(cell) if cell else ""
                        mark_found = any(mark in cell_str for mark in ["✓", "✔", "☑", "[x]", "(selected)", "√"])
                        if mark_found:
                            processed_cell = f"{cell_str} [TICKED]"
                        else:
                            processed_cell = cell_str
                        processed_row.append(processed_cell)
                    processed_rows.append(processed_row)
                # Convert entire table to a markdown-like string
                table_str = f"[Page {page_num} Table]\n"
                for row in processed_rows:
                    row_str = " | ".join(row)
                    table_str += row_str + "\n"
                tables.append(table_str)
    return tables

# Extract text from DOCX files
def extract_text_from_docx(path):
    """Extract text from paragraphs in a DOCX file."""
    doc = Document(path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

# Extract text from Excel files
def extract_text_from_excel(path):
    """Extract text from all rows in all sheets of an Excel file."""
    dfs = pd.read_excel(path, sheet_name=None)
    chunks = []
    for _, df in dfs.items():
        for row in df.itertuples(index=False):
            chunks.append(" | ".join(str(cell) for cell in row if pd.notna(cell)))
    return chunks

# Extract text from HTML files
def extract_text_from_html(path):
    """Extract visible text from an HTML file."""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        # Remove non-visible tags
        for tag in soup(['script', 'style', 'head', 'title', 'meta']):
            tag.decompose()
        # Get visible text
        lines = soup.get_text(separator='\n').splitlines()
        return [line.strip() for line in lines if line.strip()]
    except Exception as e:
        logging.warning(f"Failed to parse HTML file {path}: {e}")
        return []

# Load and extract text from all supported documents in a folder
def load_documents(folder):
    """Load and extract text from all supported documents in a folder."""
    chunks = []
    for file in glob.glob(os.path.join(folder, "*")):
        ext = os.path.splitext(file)[1].lower()
        try:
            if ext == ".pdf":
                chunks += extract_text_pages_from_pdf(file)
            elif ext in [".docx", ".doc"]:
                chunks += extract_text_from_docx(file)
            elif ext in [".xls", ".xlsx"]:
                chunks += extract_text_from_excel(file)
            elif ext in [".html", ".htm"]:
                chunks += extract_text_from_html(file)
        except Exception as e:
            logging.warning(f"Error reading {file}: {e}")
    return chunks

# =============================== Gemini Extraction ===============================
def extract_answers(rules_df, document_chunks):
    """Extract answers from Gemini by splitting document into chunks and aggregating responses."""
    global current_key_index, model
    results = {}  # final output: {normalized_title: {"Answer": ..., "Reference": ...}}

    # 1. Prepare question block
    question_blocks = []
    for idx, row in rules_df.iterrows():
        question = str(row.get("Question", "")).strip()
        title = str(row.get("Title", "")).strip()
        if not question or not title or title.lower() == "nan":
            continue
        line = f"{idx+1}. [{title}] - {question}"
        # Add metadata
        learning = str(row.get("Learning", "")).strip()
        if learning and learning.lower() != "nan":
            line += f" (Info: {learning})"
        validation_rule = str(row.get("Validation Rule", "")).strip()
        if validation_rule and validation_rule.lower() != "nan":
            line += f" (Validation: {validation_rule})"
        answer_type = str(row.get("Answer Type", "")).strip()
        if answer_type and answer_type.lower() != "nan":
            line += f" (Expected Type: {answer_type})"
        expected_format = str(row.get("Expected Form", "")).strip()
        if expected_format and expected_format.lower() != "nan":
            line += f" (Format: {expected_format})"
        keyword_columns = rules_df.columns[6:18]
        keywords = [str(row[col]).strip() for col in keyword_columns if pd.notna(row[col]) and str(row[col]).strip()]
        if keywords:
            line += f" (Keywords: {', '.join(keywords)})"
        unit_columns = rules_df.columns[18:25]
        units = [str(row[col]).strip() for col in unit_columns if pd.notna(row[col]) and str(row[col]).strip()]
        if units:
            line += f" (Units: {', '.join(units)})"
        sample = str(row.get("Sample Answer", "")).strip()
        if sample and sample.lower() != "nan":
            line += f" (Sample: {sample})"
        question_blocks.append(line)

    # 2. Combine document text and split into ~30,000-character chunks
    full_text = "\n\n---\n\n".join(document_chunks)
    chunk_size = 30000
    chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]

    logging.info(f"Total chunks created: {len(chunks)}")

    # 3. Collect raw answers from each chunk
    all_responses = []
    for idx, chunk_text in enumerate(chunks):
        prompt = f"""
You are a highly skilled tender document analyst.

Your task is to extract precise answers to the following questions from the tender document content provided below.

Respond ONLY in this strict format:
[Title]: <Answer> || <Page Number or short reference>

Guidelines:
- Base your answers only on the provided tender content. Do not guess.
- Use exact text from the document wherever possible. Avoid paraphrasing unless necessary.
- When extracting data from tables or tabular blocks:
  • Table content is marked using [Table] ... [/Table] in the document.
  • Each row in a [Table] usually represents a field like "Main Hoist", "Class of Duty", or "Remote Control".
  • Tick marks like ✓, ✔, ☑, or OCR-converted "v", "Y", "V" are replaced with [TICKED].
  • If a value is annotated as [TICKED], treat it as the **only valid option**.
  • In tables like “Main Hoist: 5.0 [TICKED], 6.0, 7.0”, always pick “5.0”.
  • Even if other values appear in context, **only select the [TICKED] one** as the final answer.
  • If multiple values are marked [TICKED], return all valid values as the answer.
- When extracting data from tables or tabular blocks, and there is only one value in the table, then return the value as the answer.
- If a value is marked as “NA” or “Not Applicable”, return “NA” as the valid answer.
- Include the exact page number or a short reference phrase from the document that supports your answer (after the `||`).
- If no valid information is found, write: [Title]: Not Found || -

Be concise. Ensure answers are specific, accurate, and reflect the source formatting.

--- DOCUMENT PART START ---
{chunk_text}
--- DOCUMENT PART END ---

Questions:
{chr(10).join(question_blocks)}
"""
        attempts = 0
        max_attempts = len(GEMINI_API_KEYS)
        while attempts < max_attempts:
            try:
                response = model.generate_content(prompt)
                response_text = response.text.strip()
                if not response_text:
                    logging.warning(f"Empty response from Gemini for chunk {idx+1}.")
                    break
                logging.info(f"Gemini response received for chunk {idx+1}.")
                all_responses.append(response_text)
                break
            except Exception as e:
                logging.error(f"Gemini call failed for chunk {idx+1}: {e}")
                attempts += 1
                if attempts < max_attempts:
                    current_key_index = (current_key_index + 1) % len(GEMINI_API_KEYS)
                    model = configure_gemini_api(current_key_index)
                else:
                    logging.error("All Gemini keys failed.")
                    break

    # 4. Merge answers from all responses
    from collections import defaultdict
    merged_answers = defaultdict(list)  # {title: [list of candidate answers from different chunks]}
    for response_text in all_responses:
        for line in response_text.splitlines():
            if ":" not in line or "||" not in line:
                continue
            try:
                title_part, rest = line.split(":", 1)
                answer, reference = rest.split("||", 1)
                title = normalize_title(title_part)
                merged_answers[title].append({
                    "Answer": answer.strip(),
                    "Reference": reference.strip()
                })
            except Exception:
                continue
    # Verify Answers
    for _, row in rules_df.iterrows():
        title = str(row.get("Title", "")).strip()
        norm_title = normalize_title(title)
        if not title or norm_title not in merged_answers:
            results[norm_title] = {"Answer": "", "Reference": ""}
            continue
        best_score = -1
        best_answer = {"Answer": "", "Reference": ""}
        for candidate in merged_answers[norm_title]:
            candidate_score = score_answer(candidate["Answer"], candidate["Reference"], row)
            if candidate_score > best_score:
                best_score = candidate_score
                best_answer = candidate
        results[norm_title] = best_answer
    return results

# Score Answer
def score_answer(answer, reference, row):
    score = 0
    raw = answer.lower()
    # 1. Skip terms penalty
    skip_terms = [
        "not found", "not available", "not in the document", "na", "n/a", "nil", "none", "not specified", "condition not met"
    ]
    if any(term in raw for term in skip_terms):
        return -1  # skip this answer
    # 2. Tick mark or checkbox logic
    tick_keywords = ["✓", "✔", "[x]", "(selected)", "☑"]
    if any(tick in answer for tick in tick_keywords):
        score += 20
    # 3. Validation Rule (regex match)
    validation = str(row.get("Validation Rule", "")).strip()
    if validation and validation.lower() != "nan":
        try:
            if re.search(validation, answer):
                score += 25
        except:
            pass
    # 4. Answer Format check
    fmt = str(row.get("Expected Form", "")).strip().lower()
    if fmt == "integer" and re.fullmatch(r"\d+", answer):
        score += 10
    elif fmt == "decimal" and re.fullmatch(r"\d+\.\d+", answer):
        score += 10
    elif fmt == "currency" and re.search(r"(rs|inr|\$|€|₹)\s?\d+", answer.lower()):
        score += 10
    elif fmt == "yes/no" and answer.strip().lower() in ["yes", "no"]:
        score += 10
    elif fmt == "date" and re.search(r"\d{2,4}[-/]\d{1,2}[-/]\d{1,4}", answer):
        score += 10
    elif fmt == "text" and len(answer.strip()) > 0:
        score += 5
    # 5. Sample Answer similarity
    sample = str(row.get("Sample Answer", "")).strip().lower()
    if sample and sample != "nan":
        similarity = difflib.SequenceMatcher(None, sample, raw).ratio()
        if similarity > 0.9:
            score += 15
        elif similarity > 0.7:
            score += 10
        elif similarity > 0.5:
            score += 5
    # 6. Answer Type matching
    atype = str(row.get("Answer Type", "")).strip().lower()
    if atype == "numeric value" and re.search(r"\d+", answer):
        score += 10
    elif atype == "text" and len(answer) > 0:
        score += 5
    elif atype == "yes/no" and answer.lower() in ["yes", "no"]:
        score += 10
    elif atype == "keyword match":
        keyword_columns = row.index[6:18]
        keywords = [str(row[col]).strip().lower() for col in keyword_columns if pd.notna(row[col])]
        if any(kw in raw for kw in keywords):
            score += 10
    # 7. Bonus for containing "[TICKED]"
    if "[ticked]" in raw:
        score += 50
    return score

# Clean unicode & Remove Ascii Characters
def clean_unicode(val):
    """Replace non-ASCII characters with '?' to avoid encoding errors."""
    if isinstance(val, str):
        return val.encode("ascii", "replace").decode("ascii")
    return val

# =============================== Main Function ===============================

def main(folder_path):
    """Main function to extract answers and write to Excel."""
    global model
    try:
        rules_df = pd.read_excel(INPUT_RULES_FILE, sheet_name="master")
        chunks = load_documents(folder_path)
        # Write extracted chunks to a DOCX file for inspection
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Extracted Document Chunks', 0)
            for chunk in chunks:
                doc.add_paragraph(chunk)
            docx_output_path = os.path.join(folder_path, 'Extracted_Text_Inspection.docx')
            doc.save(docx_output_path)
            logging.info(f"Extracted text saved for inspection: {docx_output_path}")
        except Exception as e:
            logging.error(f"Failed to write extracted text to DOCX: {e}")
        if not chunks:
            raise ValueError("No documents loaded. Ensure the folder contains valid files.")
        extracted_answers = extract_answers(rules_df, chunks)
        seen_titles = set()
        output_rows = []
        for _, row in rules_df.iterrows():
            title = str(row.get("Title", "")).strip()
            if not title or title.lower() == "nan" or title in seen_titles:
                continue
            seen_titles.add(title)
            norm_title = normalize_title(title)
            match = extracted_answers.get(norm_title)
            if match:
                ans = match["Answer"].strip().lower()
                skip_terms = [
                    "not found", "not available", "not in the document", "na",
                    "n/a", "nil", "none", "not specified", "condition not met"
                ]
                if any(term in ans for term in skip_terms):
                    answer, reference = "", ""
                else:
                    answer, reference = match["Answer"], match["Reference"]
            else:
                answer, reference = "", ""
            output_rows.append({
                "Title": title,
                "Answer": answer,
                "Reference": reference
            })
        final_df = pd.DataFrame(output_rows)
        final_df = final_df.applymap(clean_unicode)
        tender_number = ""
        tender_key = normalize_title("Tender Id")
        if tender_key in extracted_answers:
            tender_number = extracted_answers[tender_key]["Answer"].strip()
        safe_tender_number = re.sub(r'[\\/*?:"<>|]', "_", tender_number)
        if safe_tender_number:
            output_filename = os.path.join(folder_path, f"00 Summary of Tender - {safe_tender_number}.xlsx")
        else:
            output_filename = os.path.join(folder_path, "00 Extracted_Answers.xlsx")
        with pd.ExcelWriter(output_filename, engine="openpyxl") as writer:
            final_df.to_excel(writer, sheet_name="output", index=False)
        # Apply formatting
        import openpyxl
        from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
        wb = openpyxl.load_workbook(output_filename)
        ws = wb["output"]
        # Column width
        for col in ['A', 'B', 'C']:
            ws.column_dimensions[col].width = 40
        # Border
        thin = Side(border_style="thin", color="000000")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)
        # Header styling
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        # Cell formatting
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=3):
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        # Filters
        ws.auto_filter.ref = f"A1:C1"
        wb.save(output_filename)
        if os.path.exists(output_filename):
            logging.info(f"All done! File saved to: {output_filename}")
        else:
            logging.error("Something went wrong, file not saved.")
    except Exception as e:
        logging.error(f"Fatal error: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        main(sys.argv[1])
    else:
        logging.error("Folder path not provided.")